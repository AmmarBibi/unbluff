"""Deliverable -> text + citation extraction for the consistency-audit skill.

Format-agnostic by design: md/txt/tex/rst are read directly; docx via python-docx
or a stdlib zip+xml fallback; pdf via `pdftotext` / PyMuPDF / pdfminer when any is
available. This module does the MECHANICAL half only - pulling numbers, figure
references, figure embeds and captions out of the prose. Whether a claim is
actually supported, or an interpretation is consistent across sections, is the
reasoning half and belongs to Claude, not to this script.

Every extractor is a pure function over text (except the format readers, which
touch the filesystem). Nothing here decides "this is wrong" - it surfaces
candidates and tags their context so the audit can rank them.
"""
from __future__ import annotations

import os
import re
import subprocess
import zipfile
from dataclasses import dataclass
from typing import Dict, List, Optional, Tuple
from xml.etree import ElementTree as ET


class ExtractError(RuntimeError):
    """Raised when a deliverable cannot be turned into text by any available means."""


# --------------------------------------------------------------------------- #
# Deliverable -> plain text
# --------------------------------------------------------------------------- #

TEXT_EXTS = {".md", ".markdown", ".txt", ".tex", ".rst", ".rmd", ".qmd", ".org"}


def deliverable_to_text(path: str) -> str:
    """Return the plain text of a deliverable, dispatching on extension.

    Raises ExtractError with actionable guidance when a binary format cannot be
    read with the tools available on this machine.
    """
    ext = os.path.splitext(path)[1].lower()
    if ext in TEXT_EXTS:
        return _read_text(path)
    if ext == ".docx":
        return _docx_to_text(path)
    if ext == ".pdf":
        return _pdf_to_text(path)
    # Unknown extension: try as text, else give up with guidance.
    try:
        return _read_text(path)
    except (OSError, UnicodeError) as exc:
        raise ExtractError(
            "Unsupported deliverable format %r (%s). Convert it to markdown/text "
            "(e.g. `pandoc in%s -o out.md`) and re-run against the .md." % (ext, exc, ext)
        )


def _read_text(path: str) -> str:
    with open(path, encoding="utf-8", errors="replace") as handle:
        return handle.read()


def _docx_to_text(path: str) -> str:
    """docx -> text. Prefer python-docx; fall back to unzipping document.xml."""
    try:  # best fidelity if the user has python-docx
        import docx  # type: ignore
    except ImportError:
        return _docx_to_text_stdlib(path)
    try:
        document = docx.Document(path)
    except Exception as exc:  # corrupt file etc. - fall back before giving up
        try:
            return _docx_to_text_stdlib(path)
        except Exception:
            raise ExtractError("Could not read docx %r: %s" % (path, exc))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)


_WORD_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def _docx_to_text_stdlib(path: str) -> str:
    """Zero-dependency docx reader: pull text runs out of word/document.xml."""
    try:
        with zipfile.ZipFile(path) as zf:
            xml = zf.read("word/document.xml")
    except (OSError, KeyError, zipfile.BadZipFile) as exc:
        raise ExtractError("Could not open docx %r as a zip: %s" % (path, exc))
    try:
        root = ET.fromstring(xml)
    except ET.ParseError as exc:
        raise ExtractError("Malformed docx XML in %r: %s" % (path, exc))
    lines: List[str] = []
    for para in root.iter(_WORD_NS + "p"):
        text = "".join(node.text or "" for node in para.iter(_WORD_NS + "t"))
        lines.append(text)
    return "\n".join(lines)


def _pdf_to_text(path: str) -> str:
    """pdf -> text via pdftotext, then PyMuPDF, then pdfminer; else guidance."""
    # 1) poppler's pdftotext on PATH (fast, high quality, -layout keeps tables).
    try:
        proc = subprocess.run(
            ["pdftotext", "-layout", path, "-"],
            capture_output=True, timeout=120, stdin=subprocess.DEVNULL,
        )
        if proc.returncode == 0 and proc.stdout.strip():
            return proc.stdout.decode("utf-8", errors="replace")
    except (OSError, subprocess.SubprocessError):
        pass
    # 2) PyMuPDF (fitz).
    try:
        import fitz  # type: ignore
        with fitz.open(path) as doc:
            return "\n".join(page.get_text() for page in doc)
    except Exception:
        pass
    # 3) pdfminer.six.
    try:
        from pdfminer.high_level import extract_text  # type: ignore
        text = extract_text(path)
        if text and text.strip():
            return text
    except Exception:
        pass
    raise ExtractError(
        "No PDF text extractor available. Install poppler (`pdftotext`) or "
        "`pip install pymupdf` / `pdfminer.six`, or convert the PDF to text first."
    )


# --------------------------------------------------------------------------- #
# Number extraction
# --------------------------------------------------------------------------- #

@dataclass
class NumberCite:
    value: float
    raw: str
    line: int
    context: str
    decimals: int
    is_percent: bool
    unit: Optional[str]
    ref_context: bool     # preceded by "Figure/Table/Eq/Section/[..]" etc.
    looks_like_year: bool


# A number: optional sign, thousands-grouped or plain integer part, optional
# fraction, optional exponent. Percent / unit are matched separately after.
_NUMBER_RE = re.compile(
    r"(?<![\w.])"
    r"(?P<sign>[-+]?)"
    r"(?P<intpart>\d{1,3}(?:,\d{3})+|\d+)"
    r"(?P<frac>\.\d+)?"
    r"(?P<exp>[eE][-+]?\d+)?"
)

# Units we recognise immediately after a number (optionally one space).
_UNIT_RE = re.compile(
    r"\s?("
    r"%|dB|Hz|kHz|MHz|GHz|rad/s|rad|deg|°|"
    r"ms|us|µs|ns|s\b|min\b|hrs?\b|"
    r"mm|cm|km|nm|µm|um|m\b|"
    r"kg|mg|g\b|"
    r"kN|mN|N·m|Nm|N\b|"
    r"kPa|MPa|Pa|"
    r"kW|mW|W\b|kV|mV|V\b|mA|A\b|"
    r"x\b|X\b"
    r")"
)

# Words that, immediately before a number, mark it as a cross-reference /
# ordinal rather than a measured quantity.
_REF_PREFIX_RE = re.compile(
    r"(?:figure|fig|table|tbl|section|sect|sec|equation|eqn|eq|chapter|chap|ch|"
    r"appendix|app|reference|ref|step|stage|part|phase|question|q|item|version|"
    r"v|no|number|num|line|page|pg|p|slide|footnote|note|clause)\.?\s*$",
    re.IGNORECASE,
)


def _line_of(text: str, pos: int) -> int:
    return text.count("\n", 0, pos) + 1


def _is_ref_context(prefix: str) -> bool:
    """True if the text just before a number marks it as a reference/ordinal."""
    tail = prefix[-24:]
    if tail.rstrip().endswith(("[", "#")):
        return True
    return bool(_REF_PREFIX_RE.search(tail))


def _clean_float(sign: str, intpart: str, frac: str, exp: str) -> float:
    return float("%s%s%s%s" % (sign, intpart.replace(",", ""), frac, exp))


def find_numbers(text: str, context_chars: int = 60) -> List[NumberCite]:
    """Extract every numeric literal with its context and classifying tags."""
    cites: List[NumberCite] = []
    for match in _NUMBER_RE.finditer(text):
        intpart = match.group("intpart")
        frac = match.group("frac") or ""
        exp = match.group("exp") or ""
        sign = match.group("sign") or ""
        try:
            value = _clean_float(sign, intpart, frac, exp)
        except ValueError:
            continue
        start, end = match.start(), match.end()
        unit_match = _UNIT_RE.match(text, end)
        unit = unit_match.group(1) if unit_match else None
        raw = text[start:(unit_match.end() if unit_match else end)]
        is_percent = raw.rstrip().endswith("%") or unit == "%"
        decimals = len(frac) - 1 if frac else 0
        prefix = text[max(0, start - 24):start]
        ref_ctx = _is_ref_context(prefix)
        looks_year = (decimals == 0 and unit is None and not is_percent
                      and 1900 <= value <= 2099 and "," not in intpart)
        ctx_start = max(0, start - context_chars)
        ctx_end = min(len(text), end + context_chars)
        context = " ".join(text[ctx_start:ctx_end].split())
        cites.append(NumberCite(
            value=value, raw=raw, line=_line_of(text, start), context=context,
            decimals=decimals, is_percent=is_percent, unit=unit,
            ref_context=ref_ctx, looks_like_year=looks_year,
        ))
    return cites


# --------------------------------------------------------------------------- #
# Figure / table references, embeds, captions - and unfilled placeholders
# --------------------------------------------------------------------------- #

_FIG_REF_RE = re.compile(r"(?:figure|fig)\.?\s*([0-9]+(?:\.[0-9]+)?)", re.IGNORECASE)
_TBL_REF_RE = re.compile(r"(?:table|tbl)\.?\s*([0-9]+(?:\.[0-9]+)?)", re.IGNORECASE)
_FIG_CAPTION_RE = re.compile(
    r"^\s*(?:figure|fig)\.?\s*([0-9]+(?:\.[0-9]+)?)\s*[:.\-–]\s*(.+?)\s*$", re.IGNORECASE)
_TBL_CAPTION_RE = re.compile(
    r"^\s*(?:table|tbl)\.?\s*([0-9]+(?:\.[0-9]+)?)\s*[:.\-–]\s*(.+?)\s*$", re.IGNORECASE)
_MD_IMG_RE = re.compile(r"!\[[^\]]*\]\(([^)\s]+)")
_HTML_IMG_RE = re.compile(r"<img[^>]+src=[\"']([^\"']+)", re.IGNORECASE)
_TEX_IMG_RE = re.compile(r"\\includegraphics(?:\[[^\]]*\])?\{([^}]+)\}")

# A RENDERED table: a markdown separator row (|---|---|), an HTML <table>, or a
# LaTeX tabular/table environment. Used to tell a real table from a placeholder.
_MD_TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(?:\|\s*:?-{3,}:?\s*)+\|?\s*$",
                              re.MULTILINE)
_HTML_TABLE_RE = re.compile(r"<table[\s>]", re.IGNORECASE)
_TEX_TABLE_RE = re.compile(r"\\begin\{(?:tabular|longtable|table)\}", re.IGNORECASE)

# Unfilled placeholders left in a deliverable: bracketed placeholder vocabulary,
# empty / ellipsis brackets, angle-bracket placeholders, and the classic bare
# markers (TKTK, XXXX, TBD). Deliberately NOT matching [12]-style citations or
# [N/A] values.
_PLACEHOLDER_VOCAB = (
    r"table|tbl|figure|fig|graph|chart|plot|todo|tbd|fixme|placeholder|insert|"
    r"fill[\s-]?in|to[\s-]?do|to[\s-]?come|pending|value|values|data|number|"
    r"result|results|citation|cite|ref|xx+|tk"
)
_PLACEHOLDER_RE = re.compile(
    r"\[\s*\]"                                                 # empty [ ]
    r"|\[\s*\.\.\.\s*\]"                                       # [ ... ]
    r"|\[[^\]]*\b(?:" + _PLACEHOLDER_VOCAB + r")\b[^\]]*\]"    # [ ... TABLE ... ]
    r"|<[^>]*\b(?:placeholder|insert|todo|tbd)\b[^>]*>"        # <placeholder ...>
    r"|\bTKTK\b|\bXXXX+\b|\bTBD\b",                            # bare classic markers
    re.IGNORECASE,
)


def find_figure_refs(text: str) -> List[Tuple[str, int]]:
    """All in-text 'Figure N' mentions as (number, line)."""
    return [(m.group(1), _line_of(text, m.start())) for m in _FIG_REF_RE.finditer(text)]


def find_table_refs(text: str) -> List[Tuple[str, int]]:
    """All in-text 'Table N' mentions as (number, line)."""
    return [(m.group(1), _line_of(text, m.start())) for m in _TBL_REF_RE.finditer(text)]


def _captions(text: str, regex) -> Dict[str, Tuple[int, str]]:
    caps: Dict[str, Tuple[int, str]] = {}
    for lineno, line in enumerate(text.splitlines(), 1):
        m = regex.match(line)
        if m:
            caps.setdefault(m.group(1), (lineno, m.group(2).strip()))
    return caps


def find_figure_captions(text: str) -> Dict[str, Tuple[int, str]]:
    """Map 'Figure N' -> (line, caption text) for figure caption lines."""
    return _captions(text, _FIG_CAPTION_RE)


def find_table_captions(text: str) -> Dict[str, Tuple[int, str]]:
    """Map 'Table N' -> (line, caption text) for table caption lines."""
    return _captions(text, _TBL_CAPTION_RE)


def count_table_structures(text: str) -> int:
    """Number of RENDERED tables (markdown/html/latex) present in the text."""
    return (len(_MD_TABLE_SEP_RE.findall(text))
            + len(_HTML_TABLE_RE.findall(text))
            + len(_TEX_TABLE_RE.findall(text)))


def find_placeholders(text: str, cap: int = 40) -> List[Tuple[int, str]]:
    """Unfilled placeholders left in the deliverable as (line, matched text).

    Skips markdown links `[label](url)` - the label is not a placeholder.
    """
    out: List[Tuple[int, str]] = []
    for m in _PLACEHOLDER_RE.finditer(text):
        if m.group(0).startswith("[") and m.end() < len(text) and text[m.end()] == "(":
            continue  # [label](url) markdown link, not a placeholder
        out.append((_line_of(text, m.start()), m.group(0).strip()))
        if len(out) >= cap:
            break
    return out


def find_figure_embeds(path: str, text: str) -> List[str]:
    """Embedded-image identifiers: md/html/tex paths, or docx media members."""
    if os.path.splitext(path)[1].lower() == ".docx":
        try:
            with zipfile.ZipFile(path) as zf:
                return sorted(n for n in zf.namelist()
                              if n.startswith("word/media/"))
        except (OSError, zipfile.BadZipFile):
            return []
    embeds: List[str] = []
    for regex in (_MD_IMG_RE, _HTML_IMG_RE, _TEX_IMG_RE):
        embeds.extend(m.group(1) for m in regex.finditer(text))
    return embeds
